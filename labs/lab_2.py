import pandas as pd
from docx import Document
from fpdf import FPDF

df = pd.read_csv('labs/file.csv', sep='\t')
#---------------------------------------------------------
df_sorted = df.sort_values('Cost', ascending=False)
total_sum = df_sorted['Cost'].sum()
df_sorted.to_excel('labs/report.xlsx', index=False)
#---------------------------------------------------------
doc = Document()
doc.add_heading('Financial report', 0)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Name'
hdr_cells[1].text = 'Amount'
hdr_cells[2].text = 'Cost'

for _, row in df_sorted.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(row['Name'])
    row_cells[1].text = str(row['Amount'])
    row_cells[2].text = str(row['Cost'])

doc.add_paragraph(f'Total sum: {total_sum} chubrikov.')
doc.save('labs/report.docx')
#---------------------------------------------------------
class PDF(FPDF):
    def header(self):
        try:
            self.image('labs/logo.jpg', 10, 8, 33)
        except:
            pass
        self.set_y(45)
        
        self.set_font('Arial', 'B', 15)
        self.cell(0, 10, 'Financial Report', 0, 1, 'C')
        self.ln(5)
        
    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

pdf = PDF()
pdf.add_page()
pdf.set_font('Arial', '', 10)

pdf.set_font('Arial', 'B', 10)
pdf.cell(60, 10, 'Name', 1)
pdf.cell(40, 10, 'Amount', 1)
pdf.cell(40, 10, 'Cost', 1)
pdf.ln()

pdf.set_font('Arial', '', 10)
for _, row in df_sorted.iterrows():
    pdf.cell(60, 10, str(row['Name']), 1)
    pdf.cell(40, 10, str(row['Amount']), 1)
    pdf.cell(40, 10, str(row['Cost']), 1)
    pdf.ln()

pdf.ln(10)
pdf.set_font('Arial', 'B', 12)
pdf.cell(0, 10, f'Total sum: {total_sum} chubrikov.', 0, 1)

pdf.output('labs/report.pdf')
#---------------------------------------------------------
print("Отчеты успешно созданы: report.xlsx, report.docx, report.pdf")